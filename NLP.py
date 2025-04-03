#import fitz  # PyMuPDF
import pytesseract
from PIL import Image
import pandas as pd
from docx import Document
import nltk
from nltk.tokenize import word_tokenize, sent_tokenize
import os

# Set the path for Tesseract OCR executable
pytesseract.pytesseract.tesseract_cmd = r'path_to_tesseract_executable'

# Download NLTK data
nltk.download('punkt')

# Function to extract text from PDF
'''def extract_text_from_pdf(pdf_path):
    doc = fitz.open(pdf_path)
    text = ""
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        text += page.get_text()
    return text
'''
# Function to extract text from image
def extract_text_from_image(image_path):
    return pytesseract.image_to_string(Image.open(image_path))

# Function to extract text from txt file
def extract_text_from_txt(txt_path):
    with open(txt_path, 'r') as file:
        text = file.read()
    return text

# Function to extract data from Excel under a heading
def extract_data_from_excel(excel_path, heading):
    df = pd.read_excel(excel_path)
    data = df[df.apply(lambda row: row.astype(str).str.contains(heading).any(), axis=1)]
    return data

# Function to process text and find data under a heading
def find_data_under_heading(text, heading):
    sentences = sent_tokenize(text)
    heading_index = -1
    for i, sentence in enumerate(sentences):
        if heading.lower() in sentence.lower():
            heading_index = i
            break
    if heading_index == -1:
        return ""
    data = ""
    for sentence in sentences[heading_index+1:]:
        if sentence.strip() == "":
            break
        data += sentence + " "
    return data.strip()

# Function to create a Word document with extracted data
def create_word_document(output_path, data_dict, tables_dict):
    doc = Document()
    for file_type, data in data_dict.items():
        doc.add_heading(file_type, level=1)
        doc.add_paragraph(data)
    
    for file_type, table in tables_dict.items():
        doc.add_heading(file_type + " Tables", level=1)
        table = doc.add_table(rows=1, cols=len(table.columns))
        hdr_cells = table.rows[0].cells
        for i, column in enumerate(table.columns):
            hdr_cells[i].text = str(column)
        for index, row in table.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value)
    
    doc.save(output_path)

# Main function to extract data from different file types
def main():
    files = {
        'image': 'C:\Users\Athish\Downloads\3-10.jpg',
        'excel': 'C:\Users\Athish\Desktop\Book1.xlsx'
    }
    heading = "Your Heading"

    data_dict = {}
    tables_dict = {}

    for file_type, file_path in files.items():
        if file_type == 'image':
            text = extract_text_from_image(file_path)
            data = find_data_under_heading(text, heading)
            data_dict['Image'] = data
        elif file_type == 'txt':
            text = extract_text_from_txt(file_path)
            data = find_data_under_heading(text, heading)
            data_dict['Text'] = data
        elif file_type == 'excel':
            table = extract_data_from_excel(file_path, heading)
            tables_dict['Excel'] = table

    create_word_document('output.docx', data_dict, tables_dict)

if __name__ == "__main__":
    main()
